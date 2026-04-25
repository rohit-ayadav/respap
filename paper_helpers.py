"""Formatting helpers for DOCX paper generation."""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

ACCENT = RGBColor(0x1B, 0x3A, 0x5C)
ACCENT_HEX = '1B3A5C'
HEADER_BG = 'D9E2F3'
ALT_ROW = 'F2F6FC'
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))


def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color='999999', sz='4'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:color'), color)
        el.set(qn('w:space'), '0')
        borders.append(el)
    tcPr.append(borders)


def make_two_column(section):
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '360')
    section._sectPr.append(cols)


def add_page_number(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run._r.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(9)


def add_header_text(section, text):
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True
    # bottom border on header
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:color'), 'CCCCCC')
    bottom.set(qn('w:space'), '1')
    pBdr.append(bottom)
    pPr.append(pBdr)


def styled_para(doc, text, bold=False, italic=False, size=11,
                align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=6, before=0,
                font='Times New Roman', color=None, indent=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    if indent:
        p.paragraph_format.first_line_indent = Cm(indent)
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def section_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = ACCENT if level == 1 else RGBColor(0x33, 0x33, 0x33)
    if level == 1:
        pPr = h._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '8')
        bottom.set(qn('w:color'), ACCENT_HEX)
        bottom.set(qn('w:space'), '1')
        pBdr.append(bottom)
        pPr.append(pBdr)
    return h


def pro_table(doc, headers, rows, caption=None, widths=None):
    if caption:
        styled_para(doc, caption, bold=True, italic=True, size=9,
                     align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.autofit = True
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, HEADER_BG)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(9)
                r.font.color.rgb = ACCENT
    # Data
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            if ri % 2 == 1:
                set_cell_shading(cell, ALT_ROW)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(9)
    return table


def add_bullet(doc, text, bold_prefix=''):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.bold = True
        rb.font.name = 'Times New Roman'
        rb.font.size = Pt(10)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10)
    return p


def add_numbered(doc, text, bold_prefix=''):
    p = doc.add_paragraph(style='List Number')
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.bold = True
        rb.font.name = 'Times New Roman'
        rb.font.size = Pt(10)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10)
    return p


def add_code_block(doc, code_text, caption=None):
    if caption:
        styled_para(doc, caption, bold=True, italic=True, size=9, after=2)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(4)
    # Light gray background via shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F5F5F5')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)
    # Border
    pBdr = OxmlElement('w:pBdr')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), 'CCCCCC')
        el.set(qn('w:space'), '2')
        pBdr.append(el)
    pPr.append(pBdr)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def add_image(doc, path, width_inches=3.2, caption=None):
    if not os.path.exists(path):
        styled_para(doc, f'[Image not found: {os.path.basename(path)}]',
                     italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_inches))
    if caption:
        styled_para(doc, caption, italic=True, size=9,
                     align=WD_ALIGN_PARAGRAPH.CENTER, after=8)


def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), 'DDDDDD')
    bottom.set(qn('w:space'), '1')
    pBdr.append(bottom)
    pPr.append(pBdr)


def init_document():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(1.91)
        section.right_margin = Cm(1.91)
        section.different_first_page_header_footer = False
    make_two_column(doc.sections[0])
    add_page_number(doc.sections[0])
    add_header_text(doc.sections[0],
        'SafeHer: A Proactive AI-Driven Safety System for Women')
    return doc
